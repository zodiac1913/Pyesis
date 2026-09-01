from __future__ import annotations

from collections import defaultdict
from dataclasses import dataclass
from datetime import datetime, timedelta
from pathlib import Path
import re

from docx import Document
from docx.shared import Inches

from pyesis.config import AppConfig, EntryRecord
from pyesis.git_monitor import summarize_file_changes


DAY_ORDER = [
    "Monday",
    "Tuesday",
    "Wednesday",
    "Thursday",
    "Friday",
    "Saturday",
    "Sunday",
]
LIST_BULLET_LEVEL_THREE = "List Bullet 3"
LIST_BULLET_STYLE = "List Bullet"
LIST_NUMBER_STYLE = "List Number"
AI_WEEKLY_DAY_INDENT = Inches(0.25)
AI_WEEKLY_REPO_INDENT = Inches(0.5)
AI_WEEKLY_ITEM_INDENT = Inches(0.75)


@dataclass(frozen=True)
class RenderedTextChunk:
    text: str
    tags: tuple[str, ...] = ()


def _week_end_day_index(week_end_day: str) -> int:
    try:
        return DAY_ORDER.index(week_end_day)
    except ValueError:
        return DAY_ORDER.index("Thursday")


def _active_week_start(week_end_day: str, now: datetime | None = None) -> datetime:
    reference = now or datetime.now()
    end_index = _week_end_day_index(week_end_day)
    week_end = (reference + timedelta(days=(end_index - reference.weekday()) % 7)).replace(
        hour=0,
        minute=0,
        second=0,
        microsecond=0,
    )
    return week_end - timedelta(days=6)


def _week_end_date(week_start: datetime) -> datetime:
    return week_start + timedelta(days=6)


def _group_entries(entries: list[EntryRecord]) -> dict[str, dict[str, list[EntryRecord]]]:
    grouped: dict[str, dict[str, list[EntryRecord]]] = defaultdict(lambda: defaultdict(list))
    for entry in sorted(entries, key=lambda item: item.created_at):
        grouped[entry.week_start_iso][entry.day_name].append(entry)
    return grouped


def _active_week_entries(
    entries: list[EntryRecord],
    week_end_day: str,
    now: datetime | None = None,
) -> tuple[str, dict[str, list[EntryRecord]]]:
    active_week_start_iso = _active_week_start(week_end_day, now).isoformat()
    grouped = _group_entries(entries)
    return active_week_start_iso, grouped.get(active_week_start_iso, {})


def _entries_for_week_start(
    entries: list[EntryRecord],
    week_start_iso: str,
) -> dict[str, list[EntryRecord]]:
    grouped = _group_entries(entries)
    return grouped.get(week_start_iso, {})


def _group_entries_by_repo(entries: list[EntryRecord]) -> dict[str, list[EntryRecord]]:
    by_repo: dict[str, list[EntryRecord]] = defaultdict(list)
    for entry in entries:
        by_repo[entry.repo_label].append(entry)
    return {
        repo_label: sorted(repo_entries, key=lambda item: item.created_at)
        for repo_label, repo_entries in sorted(by_repo.items(), key=lambda item: item[0].lower())
    }


def render_plain_text(config: AppConfig, now: datetime | None = None) -> str:
    return "".join(chunk.text for chunk in render_text_chunks(config, now=now)).rstrip("\n") + "\n"


def render_weekly_evidence_text(
    config: AppConfig,
    now: datetime | None = None,
    week_start_iso: str | None = None,
) -> str:
    if week_start_iso is None:
        selected_week_start_iso, selected_week_entries = _active_week_entries(config.entries, config.week_end_day, now=now)
    else:
        selected_week_start_iso = week_start_iso
        selected_week_entries = _entries_for_week_start(config.entries, selected_week_start_iso)

    week_start = datetime.fromisoformat(selected_week_start_iso)
    week_end = _week_end_date(week_start)

    lines = _weekly_evidence_header_lines(week_start, week_end)

    if not selected_week_entries:
        lines.append("No captured entries for the current week.")
        return "\n".join(lines).rstrip() + "\n"

    for day_name in DAY_ORDER:
        entries = selected_week_entries.get(day_name)
        if not entries:
            continue
        lines.extend(_weekly_evidence_day_lines(day_name, entries))

    return "\n".join(lines).rstrip() + "\n"


def _weekly_evidence_header_lines(week_start: datetime, week_end: datetime) -> list[str]:
    return [
        f"Week ending: {week_end.strftime('%Y-%m-%d')}",
        f"Week start: {week_start.strftime('%Y-%m-%d')}",
        "",
    ]


def _weekly_evidence_day_lines(day_name: str, entries: list[EntryRecord]) -> list[str]:
    lines = [f"Day: {day_name}"]
    for repo_label, repo_entries in _group_entries_by_repo(entries).items():
        lines.extend(_weekly_evidence_repo_lines(repo_label, repo_entries))
    return lines


def _weekly_evidence_repo_lines(repo_label: str, repo_entries: list[EntryRecord]) -> list[str]:
    lines = [f"Repo: {repo_label}"]
    for entry in repo_entries:
        lines.extend(_weekly_evidence_entry_lines(entry))
    lines.append("")
    return lines


def _weekly_evidence_entry_lines(entry: EntryRecord) -> list[str]:
    lines = [f"- Summary: {_summary_body_text(entry.summary, entry.diff_excerpt)}"]
    evidence = _entry_evidence_line(entry)
    if evidence:
        lines.append(f"  Evidence: {evidence}")
    for label, change_line in _change_detail_lines(entry.diff_excerpt):
        lines.append(f"  {label}: {change_line}")
    warning = entry.summary_warning.strip()
    if warning:
        lines.append(f"  Warning: {warning}")
    return lines


def render_text_chunks(
    config: AppConfig,
    entry_tag_resolver=None,
    warning_comment_resolver=None,
    delete_tag_resolver=None,
    now: datetime | None = None,
) -> list[RenderedTextChunk]:
    active_week_start_iso, active_week_entries = _active_week_entries(config.entries, config.week_end_day, now=now)
    chunks: list[RenderedTextChunk] = []

    _append_week_header(chunks, active_week_start_iso)
    if active_week_entries:
        _append_week_entries(chunks, active_week_entries, entry_tag_resolver, warning_comment_resolver, delete_tag_resolver)

    return chunks


def _append_week_header(chunks: list[RenderedTextChunk], week_start_iso: str) -> None:
    week_start = datetime.fromisoformat(week_start_iso)
    week_end = _week_end_date(week_start)
    chunks.extend(RenderedTextChunk("\n") for _ in range(6))
    chunks.append(RenderedTextChunk(f"({week_end.strftime('%Y %b %d')})\n"))
    chunks.append(RenderedTextChunk("What I worked on for this week:\n"))
    chunks.append(RenderedTextChunk("\n"))


def _append_week_entries(
    chunks: list[RenderedTextChunk],
    day_map: dict[str, list[EntryRecord]],
    entry_tag_resolver,
    warning_comment_resolver,
    delete_tag_resolver,
) -> None:
    for day_name in DAY_ORDER:
        entries = day_map.get(day_name)
        if not entries:
            continue
        chunks.append(RenderedTextChunk(f"@{day_name}\n", tags=("day-heading",)))
        _append_day_repo_entries(chunks, entries, entry_tag_resolver, warning_comment_resolver, delete_tag_resolver)
        chunks.append(RenderedTextChunk("\n"))


def _append_day_repo_entries(
    chunks: list[RenderedTextChunk],
    entries: list[EntryRecord],
    entry_tag_resolver,
    warning_comment_resolver,
    delete_tag_resolver,
) -> None:
    for repo_label, repo_entries in _group_entries_by_repo(entries).items():
        chunks.append(RenderedTextChunk(f"\t• {repo_label}:\n", tags=("repo-heading",)))
        for entry in repo_entries:
            tags = _resolved_entry_tags(entry, entry_tag_resolver)
            delete_tags = _resolved_delete_tags(entry, delete_tag_resolver)
            summary_lines = _summary_lines(_summary_body_text(entry.summary, entry.diff_excerpt))
            for line_index, line in enumerate(summary_lines):
                if line_index == 0 and delete_tags:
                    chunks.append(RenderedTextChunk(f"\t\t• {line}", tags=tags))
                    chunks.append(RenderedTextChunk("  x\n", tags=delete_tags))
                    continue
                chunks.append(RenderedTextChunk(f"\t\t• {line}\n", tags=tags))
            evidence = _entry_evidence_line(entry)
            if evidence:
                chunks.append(RenderedTextChunk(f"\t\t  Evidence: {evidence}\n", tags=tags + ("evidence",)))
            for label, change_line in _change_detail_lines(entry.diff_excerpt):
                chunks.append(RenderedTextChunk(f"\t\t  {label}: {change_line}\n", tags=tags + ("evidence",)))
            warning_comment = _resolved_warning_comment(entry, warning_comment_resolver)
            if warning_comment:
                chunks.append(RenderedTextChunk(f"\t\t  {warning_comment}\n", tags=tags + ("ai-comment",)))


def _resolved_entry_tags(entry: EntryRecord, entry_tag_resolver) -> tuple[str, ...]:
    if entry_tag_resolver is not None:
        resolved = entry_tag_resolver(entry)
        if resolved is not None:
            return tuple(resolved)
    return ("heuristic",) if _is_heuristic_entry(entry) else ()


def _resolved_delete_tags(entry: EntryRecord, delete_tag_resolver) -> tuple[str, ...]:
    if delete_tag_resolver is None:
        return ()
    resolved = delete_tag_resolver(entry)
    if resolved is None:
        return ()
    return tuple(resolved)


def _resolved_warning_comment(entry: EntryRecord, warning_comment_resolver) -> str:
    if warning_comment_resolver is None:
        return ""
    return str(warning_comment_resolver(entry) or "").strip()


def _is_heuristic_entry(entry: EntryRecord) -> bool:
    source = (entry.summary_source or "").strip().lower()
    if source:
        return source == "heuristic"
    return entry.author == "Backup"


def export_docx(config: AppConfig, output_dir: Path, file_name: str | None = None) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    document = Document()
    active_week_start_iso, active_week_entries = _active_week_entries(config.entries, config.week_end_day)
    _write_week_block(document, active_week_start_iso, active_week_entries)

    if file_name:
        target = output_dir / file_name
    else:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        target = output_dir / f"weekly_changes_{timestamp}.docx"
    document.save(target)
    return target


def export_ai_weekly_report_docx(
    report_text: str,
    output_dir: Path,
    week_start_iso: str,
    provider_details: str = "",
    file_name: str | None = None,
) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    document = Document()
    week_start = datetime.fromisoformat(week_start_iso)
    week_end = _week_end_date(week_start)

    document.add_heading(f"AI Weekly Report ({week_end.strftime('%Y %b %d')})", level=1)
    if provider_details.strip():
        document.add_paragraph(f"Generated with {provider_details.strip()}")

    _write_ai_weekly_report(document, report_text)

    if file_name:
        target = output_dir / file_name
    else:
        target = output_dir / f"WhatIDidThisWeek{datetime.now().strftime('%Y%m%d')}.docx"
    document.save(target)
    return target


def _write_ai_weekly_report(document: Document, report_text: str) -> None:
    paragraph_parts: list[str] = []
    current_day: str | None = None
    current_repo: str | None = None
    repo_has_body = False

    def flush_paragraph() -> None:
        nonlocal repo_has_body
        if not paragraph_parts:
            return
        paragraph = document.add_paragraph(style=LIST_BULLET_STYLE)
        paragraph.paragraph_format.left_indent = AI_WEEKLY_ITEM_INDENT
        paragraph.paragraph_format.first_line_indent = None
        paragraph.add_run(_clean_weekly_report_item_text(" ".join(paragraph_parts).strip()))
        paragraph_parts.clear()
        repo_has_body = True

    for raw_line in report_text.splitlines():
        stripped = raw_line.strip()
        if not stripped:
            flush_paragraph()
            continue

        markdown_heading = _weekly_markdown_heading(stripped)
        if markdown_heading:
            flush_paragraph()
            heading_level, heading_text = markdown_heading
            heading_text = _clean_weekly_report_inline_text(heading_text)
            day_heading = _parse_weekly_report_day_heading(heading_text)
            if day_heading:
                _add_ai_weekly_day_paragraph(document, day_heading)
                current_day = day_heading
                current_repo = None
                repo_has_body = False
                continue

            if current_day and not current_repo:
                _add_ai_weekly_repo_paragraph(document, heading_text)
                current_repo = heading_text
                repo_has_body = False
                continue

            paragraph = document.add_paragraph(heading_text)
            paragraph.paragraph_format.left_indent = AI_WEEKLY_DAY_INDENT
            del heading_level
            continue

        day_heading = _parse_weekly_report_day_heading(stripped)
        if day_heading:
            flush_paragraph()
            _add_ai_weekly_day_paragraph(document, day_heading)
            current_day = day_heading
            current_repo = None
            repo_has_body = False
            continue

        repo_heading = _parse_weekly_report_repo_heading(stripped)
        if repo_heading:
            flush_paragraph()
            _add_ai_weekly_repo_paragraph(document, repo_heading)
            current_repo = repo_heading
            repo_has_body = False
            continue

        list_item = _weekly_report_list_item(stripped)
        if list_item:
            flush_paragraph()
            style, body = list_item
            paragraph = document.add_paragraph(_clean_weekly_report_item_text(body), style=style)
            paragraph.paragraph_format.left_indent = AI_WEEKLY_ITEM_INDENT
            paragraph.paragraph_format.first_line_indent = None
            repo_has_body = True
            continue

        clean_line = _clean_weekly_report_inline_text(stripped)
        if _looks_like_weekly_report_repo_heading(clean_line, current_day, current_repo, repo_has_body):
            flush_paragraph()
            _add_ai_weekly_repo_paragraph(document, clean_line)
            current_repo = clean_line
            repo_has_body = False
            continue

        paragraph_parts.append(clean_line)

    flush_paragraph()


def _add_ai_weekly_day_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.left_indent = AI_WEEKLY_DAY_INDENT
    paragraph.paragraph_format.first_line_indent = None


def _add_ai_weekly_repo_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.left_indent = AI_WEEKLY_REPO_INDENT
    paragraph.paragraph_format.first_line_indent = None


def _weekly_markdown_heading(text: str) -> tuple[int, str] | None:
    if not text.startswith("#"):
        return None
    level = 0
    for char in text:
        if char != "#":
            break
        level += 1
    if level == 0 or level > 6:
        return None
    heading_text = text[level:].strip()
    if not heading_text:
        return None
    return level, heading_text


def _weekly_report_list_item(text: str) -> tuple[str, str] | None:
    if text.startswith(("- ", "* ")):
        return LIST_BULLET_STYLE, text[2:].strip()

    marker, separator, remainder = text.partition(".")
    if separator and marker.isdigit() and remainder.startswith(" "):
        return LIST_NUMBER_STYLE, remainder.strip()
    return None


def _parse_weekly_report_day_heading(text: str) -> str | None:
    clean_text = _clean_weekly_report_inline_text(text).strip().rstrip(":")
    if clean_text in DAY_ORDER:
        return clean_text
    if clean_text.startswith("Day:"):
        candidate = clean_text.partition(":")[2].strip().rstrip(":")
        if candidate in DAY_ORDER:
            return candidate
    if clean_text.startswith("@"):
        candidate = clean_text[1:].strip().rstrip(":")
        if candidate in DAY_ORDER:
            return candidate
    return None


def _parse_weekly_report_repo_heading(text: str) -> str | None:
    clean_text = _clean_weekly_report_inline_text(text).strip()
    if not clean_text.startswith("Repo:"):
        return None
    repo_name = clean_text.partition(":")[2].strip().rstrip(":")
    return repo_name or None


def _looks_like_weekly_report_repo_heading(
    text: str,
    current_day: str | None,
    current_repo: str | None,
    repo_has_body: bool,
) -> bool:
    if not current_day or current_repo is not None or repo_has_body:
        return False
    if not text or text in DAY_ORDER:
        return False
    if len(text) > 80:
        return False
    return not any(mark in text for mark in (". ", ": ", "! ", "? ")) and not text.endswith((".", "!", "?"))


def _clean_weekly_report_inline_text(text: str) -> str:
    clean_text = text.strip()
    clean_text = re.sub(r"^>+\s*", "", clean_text)
    clean_text = re.sub(r"\*\*(.*?)\*\*", r"\1", clean_text)
    clean_text = re.sub(r"__(.*?)__", r"\1", clean_text)
    clean_text = re.sub(r"`([^`]*)`", r"\1", clean_text)
    return clean_text.strip()


def _clean_weekly_report_item_text(text: str) -> str:
    clean_text = _clean_weekly_report_inline_text(text)
    if clean_text[:8].lower() == "summary:":
        return clean_text[8:].strip()
    return clean_text


def _write_week_block(
    document: Document,
    week_start_iso: str,
    day_map: dict[str, list[EntryRecord]],
) -> None:
    week_start = datetime.fromisoformat(week_start_iso)
    week_end = _week_end_date(week_start)
    for _ in range(6):
        document.add_paragraph("")
    document.add_paragraph(f"({week_end.strftime('%Y %b %d')})")
    document.add_paragraph("What I worked on for this week:")

    for day_name in DAY_ORDER:
        entries = day_map.get(day_name)
        if not entries:
            continue
        document.add_paragraph(f"@{day_name}")
        for repo_label, repo_entries in _group_entries_by_repo(entries).items():
            repo_paragraph = document.add_paragraph(style="List Bullet")
            repo_paragraph.paragraph_format.left_indent = None
            repo_paragraph.paragraph_format.first_line_indent = None
            repo_paragraph.add_run(f"{repo_label}:")
            _write_repo_entries(document, repo_entries)


def _write_repo_entries(document: Document, repo_entries: list[EntryRecord]) -> None:
    for entry in repo_entries:
        for idx, line in enumerate(_summary_lines(_summary_body_text(entry.summary, entry.diff_excerpt))):
            style = "List Bullet 2" if idx == 0 else LIST_BULLET_LEVEL_THREE
            entry_paragraph = document.add_paragraph(style=style)
            entry_paragraph.paragraph_format.left_indent = None
            entry_paragraph.paragraph_format.first_line_indent = None
            entry_paragraph.add_run(line)
        evidence = _entry_evidence_line(entry)
        if evidence:
            evidence_paragraph = document.add_paragraph(style=LIST_BULLET_LEVEL_THREE)
            evidence_paragraph.paragraph_format.left_indent = None
            evidence_paragraph.paragraph_format.first_line_indent = None
            evidence_paragraph.add_run(f"Evidence: {evidence}")
        for label, change_line in _change_detail_lines(entry.diff_excerpt):
            paragraph = document.add_paragraph(style=LIST_BULLET_LEVEL_THREE)
            paragraph.paragraph_format.left_indent = None
            paragraph.paragraph_format.first_line_indent = None
            paragraph.add_run(f"{label}: {change_line}")


def _summary_lines(summary: str) -> list[str]:
    lines = [line.strip().lstrip("-• ").strip() for line in summary.splitlines() if line.strip()]
    return lines or [summary.strip()]


def _entry_evidence_line(entry: EntryRecord) -> str:
    inline_evidence = _summary_inline_evidence(entry.summary)
    if inline_evidence and _evidence_has_line_number(inline_evidence) and _evidence_matches_changed_line(inline_evidence, entry.diff_excerpt):
        return inline_evidence

    changes = summarize_file_changes(entry.diff_excerpt)
    for change in changes:
        if change.added_line_samples:
            line_no, snippet = change.added_line_samples[0]
            return f"{change.path}:{line_no} \"{snippet}\""
        if change.added_samples:
            return f"{change.path} \"{change.added_samples[0]}\""
    return inline_evidence


def _summary_body_text(summary: str, diff_excerpt: str = "") -> str:
    body, _separator, _evidence = summary.partition(" Evidence: ")
    stripped = body.strip()
    resolved = stripped or summary.strip()
    if not diff_excerpt.strip():
        return resolved
    return _strip_unverified_around_claims(resolved, diff_excerpt)


def _summary_inline_evidence(summary: str) -> str:
    _body, separator, evidence = summary.partition(" Evidence: ")
    if not separator:
        return ""
    return evidence.strip().rstrip(".")


def _evidence_has_line_number(evidence: str) -> bool:
    return bool(re.search(r"^[^\s:]+:\d+\s+\"", evidence.strip(), flags=re.IGNORECASE))


def _evidence_matches_changed_line(evidence: str, diff_excerpt: str) -> bool:
    match = re.match(r'^(?P<path>[^\s:]+):(?P<line>\d+)\s+\"(?P<snippet>.*)\"$', evidence.strip())
    if not match:
        return False

    evidence_path = match.group("path").replace("\\", "/").strip().lower()
    evidence_line = int(match.group("line"))
    evidence_snippet = _normalize_evidence_snippet(match.group("snippet"))
    if not evidence_snippet:
        return False

    for change in summarize_file_changes(diff_excerpt):
        change_path = change.path.replace("\\", "/").strip().lower()
        if change_path != evidence_path:
            continue
        for line_no, snippet in change.added_line_samples:
            if line_no != evidence_line:
                continue
            if _normalize_evidence_snippet(snippet) == evidence_snippet:
                return True
    return False


def _normalize_evidence_snippet(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip())


def _strip_unverified_around_claims(summary_body: str, diff_excerpt: str) -> str:
    changed_snippets = _changed_snippet_set(diff_excerpt)
    if not changed_snippets:
        return summary_body

    marker = " around '"
    cursor = 0
    parts: list[str] = []

    while True:
        start = summary_body.find(marker, cursor)
        if start < 0:
            parts.append(summary_body[cursor:])
            break

        snippet_start = start + len(marker)
        snippet_end = summary_body.find("'", snippet_start)
        if snippet_end < 0:
            parts.append(summary_body[cursor:])
            break

        parts.append(summary_body[cursor:start])
        snippet = _normalize_evidence_snippet(summary_body[snippet_start:snippet_end])
        if snippet in changed_snippets:
            parts.append(summary_body[start:snippet_end + 1])
        cursor = snippet_end + 1

    return re.sub(r"\s+", " ", "".join(parts)).strip()


def _changed_snippet_set(diff_excerpt: str) -> set[str]:
    snippets: set[str] = set()
    for change in summarize_file_changes(diff_excerpt):
        _add_normalized_snippets(snippets, (snippet for _line_no, snippet in change.added_line_samples))
        _add_normalized_snippets(snippets, change.added_samples)
        _add_normalized_snippets(snippets, change.removed_samples)
    return snippets


def _add_normalized_snippets(target: set[str], source: list[str] | tuple[str, ...] | object) -> None:
    for snippet in source:
        normalized = _normalize_evidence_snippet(str(snippet))
        if normalized:
            target.add(normalized)


def _change_detail_lines(diff_excerpt: str) -> list[tuple[str, str]]:
    removed_lines: list[str] = []
    added_lines: list[str] = []
    added_fallback = ""
    removed_fallback = ""

    def collect_change_details() -> list[tuple[str, str]]:
        if removed_lines and added_lines:
            pair = _first_null_check_pair(removed_lines, added_lines) or _first_changed_line_pair(removed_lines, added_lines)
            if pair is not None:
                return [("Before", pair[0]), ("After", pair[1])]
        if added_lines:
            return [("Added", added_lines[0])]
        if removed_lines:
            return [("Removed", removed_lines[0])]
        return []

    def flush() -> list[tuple[str, str]]:
        nonlocal removed_lines, added_lines, added_fallback, removed_fallback
        details = collect_change_details()
        if not details:
            if added_lines and not added_fallback:
                added_fallback = added_lines[0]
            if removed_lines and not removed_fallback:
                removed_fallback = removed_lines[0]
        removed_lines = []
        added_lines = []
        return details

    for raw_line in diff_excerpt.splitlines():
        if raw_line.startswith(("@@", " ")):
            details = flush()
            if details:
                return details
            continue

        if raw_line.startswith(("diff --git ", "index ", "--- ", "+++ ", "\\")):
            continue
        if raw_line.startswith("-"):
            removed_lines.append(raw_line[1:].strip())
            continue
        if raw_line.startswith("+"):
            added_lines.append(raw_line[1:].strip())

    details = flush()
    if details:
        return details
    if added_fallback:
        return [("Added", added_fallback)]
    if removed_fallback:
        return [("Removed", removed_fallback)]
    return []


def _first_null_check_pair(removed_lines: list[str], added_lines: list[str]) -> tuple[str, str] | None:
    if not removed_lines or not added_lines:
        return None

    for before_line in removed_lines:
        if not before_line:
            continue
        for after_line in added_lines:
            if not after_line or before_line == after_line:
                continue
            if _looks_like_null_check_change(before_line, after_line):
                return before_line, after_line
    return None


def _first_changed_line_pair(removed_lines: list[str], added_lines: list[str]) -> tuple[str, str] | None:
    if not removed_lines or not added_lines:
        return None

    for before_line in removed_lines:
        if not before_line:
            continue
        for after_line in added_lines:
            if not after_line or before_line == after_line:
                continue
            return before_line, after_line
    return None


def _looks_like_null_check_change(before_line: str, after_line: str) -> bool:
    del before_line
    lowered = after_line.lower()
    null_markers = (
        "??",
        "?.",
        "== null",
        "!= null",
        " is none",
        " is not none",
        ".notempty(",
        " if ",
        " else ",
    )
    if any(marker in lowered for marker in null_markers):
        return True
    # Match C# / JS / Python ternary-like guards that often encode null fallback.
    return bool(re.search(r"\?.+:.+", after_line))